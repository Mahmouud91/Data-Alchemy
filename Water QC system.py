import tkinter as tk
from tkinter import ttk, messagebox, font, filedialog
from datetime import datetime, date, timedelta
from tkcalendar import Calendar, DateEntry
import os
import pandas as pd
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import matplotlib.dates as mdates

class WaterQCApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Pharmaceutical Water QC System")
        
        # Make window fit screen
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.root.geometry(f"{int(screen_width*0.8)}x{int(screen_height*0.8)}")
        self.root.state('zoomed')  # Start maximized
        
        # Custom colors
        self.bg_color = '#f0f8ff'  # Alice blue
        self.button_color = '#4682b4'  # Steel blue
        self.button_hover = '#5f9ea0'  # Cadet blue
        self.text_color = '#333333'
        self.accent_color = '#2e8b57'  # Sea green
        
        # Custom fonts
        self.title_font = font.Font(family='Segoe UI', size=14, weight='bold')
        self.label_font = font.Font(family='Segoe UI', size=10)
        self.button_font = font.Font(family='Segoe UI', size=10, weight='bold')
        self.table_font = font.Font(family='Segoe UI', size=9)
        self.table_header_font = font.Font(family='Segoe UI', size=9, weight='bold')
        
        # Configure styles
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Main frame style
        self.style.configure('Main.TFrame', background=self.bg_color)
        
        # Button styles
        self.style.configure('TButton', 
                           foreground='white',
                           background=self.button_color,
                           font=self.button_font,
                           padding=6,
                           borderwidth=1)
        self.style.map('TButton',
                      background=[('active', self.button_hover), ('disabled', '#d3d3d3')])
        
        # Label styles
        self.style.configure('TLabel', font=self.label_font, background=self.bg_color)
        self.style.configure('Title.TLabel', font=self.title_font, foreground=self.accent_color)
        
        # Entry styles
        self.style.configure('TEntry', font=self.label_font, padding=5)
        
        # Combobox styles
        self.style.configure('TCombobox', font=self.label_font, padding=5)
        
        # Notebook styles
        self.style.configure('TNotebook', background=self.bg_color)
        self.style.configure('TNotebook.Tab', font=self.button_font, padding=[10, 5])
        
        # Treeview styles
        self.style.configure('Treeview', 
                           font=self.table_font, 
                           rowheight=25,
                           background='white',
                           fieldbackground='white')
        self.style.configure('Treeview.Heading', 
                           font=self.table_header_font,
                           background=self.button_color,
                           foreground='white')
        self.style.map('Treeview.Heading',
                      background=[('active', self.button_hover)])
        
        # Database setup
        self.DB_FILES = {
            "Daily_Micro": "daily_microbiology.csv",
            "Daily_Chem": "daily_chemistry.csv",
            "Monthly_Micro": "monthly_microbiology.csv",
            "Monthly_Chem": "monthly_chemistry.csv",
            "Sanitization_Micro": "sanitization_microbiology.csv",
            "Sanitization_Chem": "sanitization_chemistry.csv"
        }
        self.initialize_databases()
        
        # Data storage
        self.current_data = []
        
        # Main container
        main_frame = ttk.Frame(root, style='Main.TFrame')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Title
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill='x', pady=(0, 20))
        ttk.Label(title_frame, text="Pharmaceutical Water QC System", style='Title.TLabel').pack(side='left')
        
        # Export button in top right
        export_btn = ttk.Button(title_frame, text="EXPORT TO DATABASE", command=self.export_data)
        export_btn.pack(side='right', padx=10)
        
        # Control panel frame
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=10)
        
        # Test type selection
        ttk.Label(control_frame, text="1. Select Test Type:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.test_type = ttk.Combobox(control_frame, values=["After Sanitization", "Monthly", "Daily"], state='readonly')
        self.test_type.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        self.test_type.bind("<<ComboboxSelected>>", self.update_test_ui)

        # Date selection
        ttk.Label(control_frame, text="2. Select Date:").grid(row=0, column=2, padx=5, pady=5, sticky='e')
        self.date_entry = ttk.Entry(control_frame)
        self.date_entry.grid(row=0, column=3, padx=5, pady=5, sticky='w')
        self.date_entry.insert(0, date.today().strftime("%Y-%m-%d"))
        ttk.Button(control_frame, text="📅", command=self.open_calendar, width=3).grid(row=0, column=4, padx=5)

        # Day selection (only for daily)
        self.day_label = ttk.Label(control_frame, text="3. Select Day:")
        self.day_combo = ttk.Combobox(control_frame, values=["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday"], state='readonly')
        
        # Additional point selection (only for daily)
        self.add_point_label = ttk.Label(control_frame, text="4. Add Additional Point:")
        self.add_point_combo = ttk.Combobox(control_frame, state='readonly')
        self.add_point_btn = ttk.Button(control_frame, text="Add", command=self.add_additional_point, width=5)
        
        # Notebook tabs
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill='both', expand=True, pady=10)
        
        # Setup tabs
        self.setup_micro_tab()
        self.setup_chem_tab()
        self.setup_results_viewer_tab()
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready", foreground='blue')
        self.status_label.pack(fill='x', pady=10)

        # Define all points and limits
        self.ALL_POINTS = [
            "city", "feed_water", "after_cl", "Before_sand_filter", "after_sand_filter",
            "After_Soft_1", "After_Soft_2", "After_10µFilter", "after_soft_tank", "after_smbs",
            "RO1_A", "RO1_B", "RO1_AB", "RO2", "After_EDI", "Before_PW_tank", "loop_supply",
            "loop_return", "after_heat_exchange", "UV_lamp", "PW1", "PW2", "PW3", "PWMb", "PW4", "PW5"
        ]
        
        # Microbiology limits
        self.CFU_LIMITS = {
            "city": 500, "feed_water": 500, "after_cl": 500, "Before_sand_filter": 500,
            "after_sand_filter": 500, "After_Soft_1": 500, "After_Soft_2": 500,
            "After_10µFilter": 500, "after_soft_tank": 500, "after_smbs": 500,
            "RO1_A": 500, "RO1_B": 500, "RO1_AB": 500, "RO2": 100, "After_EDI": 100,
            "Before_PW_tank": 100, "loop_supply": 100, "loop_return": 100,
            "after_heat_exchange": 100, "UV_lamp": 100, "PW1": 100, "PW2": 100,
            "PW3": 100, "PWMb": 100, "PW4": 100, "PW5": 100
        }
        
        # Chemistry limits
        self.CHEM_LIMITS = {
            "Conductivity": {
                "city": 1000, "feed_water": 1000, "after_cl": 1000, "Before_sand_filter": 1000,
                "after_sand_filter": 1000, "After_Soft_1": 1000, "After_Soft_2": 1000,
                "After_10µFilter": 1000, "after_soft_tank": 1000, "after_smbs": 1000,
                "RO1_A": 40, "RO1_B": 40, "RO1_AB": 40, "RO2": 40, "After_EDI": 1.3,
                "Before_PW_tank": 1.3, "loop_supply": 1.3, "loop_return": 1.3,
                "after_heat_exchange": 1.3, "UV_lamp": 1.3, "PW1": 1.3, "PW2": 1.3,
                "PW3": 1.3, "PWMb": 1.3, "PW4": 1.3, "PW5": 1.3
            },
            "Cl_Allowed_Points": ["city", "after_sand_filter", "After_Soft_2"]
        }
        
        # Daily test points
        self.DAILY_MICRO_POINTS = {
            "Sunday": ["PW1", "PW2", "PW3","PW4", "PW5","RO2", "loop_return", "loop_supply", "After_Soft_1"],
            "Monday": ["PW1", "PW2", "PW3","PW4", "PW5","RO2", "loop_return", "loop_supply", "After_Soft_2", "After_EDI"],
            "Tuesday": ["PW1", "PW2", "PW3","PW4", "PW5","RO2", "loop_return", "loop_supply", "after_soft_tank", "Before_PW_tank"],
            "Wednesday": ["PW1", "PW2", "PW3","PW4", "PW5","RO2", "loop_return", "loop_supply", "after_heat_exchange", "UV_lamp"],
            "Thursday": ["PW1", "PW2", "PW3","PW4", "PW5","RO2", "loop_return", "loop_supply", "PWMb"]
        }
        
        self.DAILY_CHEM_POINTS = {
            "Sunday": ["PW1", "PW2", "PW3","PW4", "PW5", "loop_return", "loop_supply", "RO2"],
            "Monday": ["PW1", "PW2", "PW3","PW4", "PW5", "loop_return", "loop_supply", "After_Soft_2", "After_EDI"],
            "Tuesday": ["PW1", "PW2", "PW3","PW4", "PW5", "loop_return", "loop_supply", "Before_PW_tank"],
            "Wednesday": ["PW1", "PW2", "PW3","PW4", "PW5", "loop_return", "loop_supply", "after_heat_exchange", "UV_lamp"],
            "Thursday": ["PW1", "PW2", "PW3","PW4", "PW5", "loop_return", "loop_supply", "PWMb"]
        }
        
        self.MONTHLY_CHEM = ["city", "Before_sand_filter", "After_Soft_2", "RO1_A", "RO1_B", "RO1_AB"]
        self.MONTHLY_MICRO = self.MONTHLY_CHEM + ["feed_water", "after_cl", "after_sand_filter", "After_10µFilter", "after_smbs"]
        
        # Initialize UI
        self.update_test_ui()

    def initialize_databases(self):
        """Create database files with headers in QC_Databases folder using pandas"""
        os.makedirs("QC_Databases", exist_ok=True)
        
        for db_key, filename in self.DB_FILES.items():
            filepath = os.path.join("QC_Databases", filename)
            if not os.path.exists(filepath):
                if "Micro" in db_key:
                    df = pd.DataFrame(columns=[
                        "Date", "Test Type", "Day", "Point",
                        "Total Count", "Coliforms", "Pseudomonas",
                        "Status", "Comments"
                    ])
                else:
                    df = pd.DataFrame(columns=[
                        "Date", "Test Type", "Day", "Point",
                        "Conductivity", "Oxidizable", "Cl Test",
                        "Status", "Comments"
                    ])
                df.to_csv(filepath, index=False)
        print("Database files initialized in QC_Databases folder")

    def setup_results_viewer_tab(self):
        """New tab for viewing historical results and generating reports"""
        self.results_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.results_frame, text="Results Viewer")
        
        # Paned window for resizable sections
        paned_window = ttk.PanedWindow(self.results_frame, orient='vertical')
        paned_window.pack(fill='both', expand=True)
        
        # Control frame
        control_frame = ttk.Frame(paned_window)
        paned_window.add(control_frame, weight=1)
        
        # Date range selection
        date_frame = ttk.Frame(control_frame)
        date_frame.pack(fill='x', pady=5)
        
        ttk.Label(date_frame, text="From:").pack(side='left', padx=5)
        self.date_from = DateEntry(date_frame, date_pattern='y-mm-dd')
        self.date_from.pack(side='left', padx=5)
        
        ttk.Label(date_frame, text="To:").pack(side='left', padx=5)
        self.date_to = DateEntry(date_frame, date_pattern='y-mm-dd')
        self.date_to.pack(side='left', padx=5)
        
        # Test type selection
        type_frame = ttk.Frame(control_frame)
        type_frame.pack(fill='x', pady=5)
        
        ttk.Label(type_frame, text="Test Type:").pack(side='left', padx=5)
        self.results_test_type = ttk.Combobox(type_frame, 
                                            values=["Daily", "After Sanitization"], 
                                            state='readonly')
        self.results_test_type.pack(side='left', padx=5)
        self.results_test_type.set("Daily")
        
        # Data type selection
        data_frame = ttk.Frame(control_frame)
        data_frame.pack(fill='x', pady=5)
        
        ttk.Label(data_frame, text="Data Type:").pack(side='left', padx=5)
        self.results_data_type = ttk.Combobox(data_frame, 
                                             values=["Microbiology", "Chemistry"], 
                                             state='readonly')
        self.results_data_type.pack(side='left', padx=5)
        self.results_data_type.set("Microbiology")
        
        # Buttons
        button_frame = ttk.Frame(control_frame)
        button_frame.pack(fill='x', pady=10)
        
        ttk.Button(button_frame, text="Load Data", command=self.load_results_data).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Generate Word Report", command=self.generate_word_report).pack(side='left', padx=5)
        
        # Results display area
        results_display_frame = ttk.Frame(paned_window)
        paned_window.add(results_display_frame, weight=3)
        
        # Notebook for table and graph views
        self.results_display_notebook = ttk.Notebook(results_display_frame)
        self.results_display_notebook.pack(fill='both', expand=True)
        
        # Table tab
        table_frame = ttk.Frame(self.results_display_notebook)
        self.results_display_notebook.add(table_frame, text="Table View")
        
        # Table with scrollbars
        tree_scroll = ttk.Frame(table_frame)
        tree_scroll.pack(fill='both', expand=True)
        
        y_scroll = ttk.Scrollbar(tree_scroll)
        y_scroll.pack(side='right', fill='y')
        
        x_scroll = ttk.Scrollbar(tree_scroll, orient='horizontal')
        x_scroll.pack(side='bottom', fill='x')
        
        self.results_table = ttk.Treeview(tree_scroll, 
                                       yscrollcommand=y_scroll.set,
                                       xscrollcommand=x_scroll.set)
        self.results_table.pack(fill='both', expand=True)
        
        y_scroll.config(command=self.results_table.yview)
        x_scroll.config(command=self.results_table.xview)
        
        # Graph tab
        graph_frame = ttk.Frame(self.results_display_notebook)
        self.results_display_notebook.add(graph_frame, text="Graph View")
        
        self.graph_canvas_frame = ttk.Frame(graph_frame)
        self.graph_canvas_frame.pack(fill='both', expand=True)

    def load_results_data(self):
        """Load historical data based on selected criteria"""
        date_from = self.date_from.get_date()
        date_to = self.date_to.get_date()
        test_type = self.results_test_type.get()
        data_type = self.results_data_type.get()
        
        if date_from > date_to:
            messagebox.showerror("Error", "End date must be after start date")
            return
        
        # Determine which file to load
        file_key = f"{'Sanitization' if test_type == 'After Sanitization' else 'Daily'}_{'Micro' if data_type == 'Microbiology' else 'Chem'}"
        filepath = os.path.join("QC_Databases", self.DB_FILES[file_key])
        
        if not os.path.exists(filepath):
            messagebox.showerror("Error", f"No data file found for {test_type} {data_type}")
            return
        
        try:
            # Load and filter data
            df = pd.read_csv(filepath)
            df['Date'] = pd.to_datetime(df['Date'])
            mask = (df['Date'] >= pd.to_datetime(date_from)) & (df['Date'] <= pd.to_datetime(date_to))
            filtered_df = df[mask]
            
            if filtered_df.empty:
                messagebox.showinfo("Info", "No data found for selected date range")
                return
            
            # Clear previous data
            for item in self.results_table.get_children():
                self.results_table.delete(item)
            
            # Configure columns based on data type
            self.results_table["columns"] = list(filtered_df.columns)
            for col in filtered_df.columns:
                self.results_table.heading(col, text=col)
                self.results_table.column(col, width=100, anchor='center')
            
            # Add data to table
            for _, row in filtered_df.iterrows():
                self.results_table.insert("", 'end', values=list(row))
            
            # Update graph
            self.update_graph(filtered_df, data_type)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load data: {str(e)}")

    def update_graph(self, df, data_type):
        """Update the graph with loaded data"""
        # Clear previous graph
        for widget in self.graph_canvas_frame.winfo_children():
            widget.destroy()
        
        # Create figure
        fig = Figure(figsize=(8, 4), dpi=100)
        ax = fig.add_subplot(111)
        
        # Convert dates for plotting
        dates = [datetime.strptime(d, '%Y-%m-%d') for d in df['Date']]
        
        if data_type == "Microbiology":
            # Plot microbiology data
            points = df['Point'].unique()
            
            for point in points:
                point_data = df[df['Point'] == point]
                counts = point_data['Total Count'].apply(lambda x: float(x) if str(x).isdigit() else 0)
                ax.plot(dates, counts, 'o-', label=point)
            
            ax.set_ylabel('CFU/mL')
            ax.set_title('Microbiology Results Over Time')
            
        else:
            # Plot chemistry data
            points = df['Point'].unique()
            
            for point in points:
                point_data = df[df['Point'] == point]
                conductivity = point_data['Conductivity'].apply(lambda x: float(x) if pd.notna(x) and str(x).replace('.', '').isdigit() else 0)
                ax.plot(dates, conductivity, 'o-', label=point)
            
            ax.set_ylabel('Conductivity (µS/cm)')
            ax.set_title('Chemistry Results Over Time')
        
        # Format x-axis
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
        ax.xaxis.set_major_locator(mdates.DayLocator(interval=max(1, len(dates)//5)))
        fig.autofmt_xdate()
        
        ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        ax.grid(True)
        
        # Embed in Tkinter
        canvas = FigureCanvasTkAgg(fig, master=self.graph_canvas_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True)

    def generate_word_report(self):
        """Generate a Word report from the loaded data"""
        if not self.results_table.get_children():
            messagebox.showerror("Error", "No data loaded to generate report")
            return
        
        # Ask for save location
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Save Report As"
        )
        
        if not filepath:
            return  # User cancelled
        
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Water QC Report', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add report details
            test_type = self.results_test_type.get()
            data_type = self.results_data_type.get()
            date_from = self.date_from.get_date().strftime('%Y-%m-%d')
            date_to = self.date_to.get_date().strftime('%Y-%m-%d')
            
            details = doc.add_paragraph()
            details.add_run(f"Test Type: {test_type}\n").bold = True
            details.add_run(f"Data Type: {data_type}\n").bold = True
            details.add_run(f"Date Range: {date_from} to {date_to}\n").bold = True
            details.add_run("\n")
            
            # Add summary statistics
            doc.add_heading('Summary Statistics', level=2)
            
            # Get data from table
            data = []
            columns = self.results_table["columns"]
            for item in self.results_table.get_children():
                values = self.results_table.item(item)['values']
                data.append(dict(zip(columns, values)))
            
            df = pd.DataFrame(data)
            
            # Add statistics table
            if data_type == "Microbiology":
                stats = df.groupby('Point')['Total Count'].agg(['count', 'mean', 'max'])
                stats.columns = ['Samples', 'Average CFU/mL', 'Max CFU/mL']
            else:
                stats = df.groupby('Point')['Conductivity'].agg(['count', 'mean', 'max'])
                stats.columns = ['Samples', 'Average Conductivity', 'Max Conductivity']
            
            # Add statistics table to document
            table = doc.add_table(stats.shape[0]+1, stats.shape[1])
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(stats.columns):
                hdr_cells[i].text = col
            
            # Data rows
            for i, (index, row) in enumerate(stats.iterrows(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = str(index)
                for j, value in enumerate(row, 1):
                    row_cells[j].text = f"{value:.2f}"
            
            # Add non-conforming results section
            doc.add_heading('Non-Conforming Results', level=2)
            non_conforming = df[df['Status'].str.contains('Non-Conform', na=False)]
            
            if len(non_conforming) > 0:
                table = doc.add_table(non_conforming.shape[0]+1, non_conforming.shape[1])
                
                # Header row
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(non_conforming.columns):
                    hdr_cells[i].text = col
                
                # Data rows
                for i, (_, row) in enumerate(non_conforming.iterrows(), 1):
                    row_cells = table.rows[i].cells
                    for j, value in enumerate(row):
                        row_cells[j].text = str(value)
            else:
                doc.add_paragraph("No non-conforming results found in this period.")
            
            # Save document
            doc.save(filepath)
            messagebox.showinfo("Success", f"Report saved successfully to:\n{filepath}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report: {str(e)}")

    def open_calendar(self):
        """Open calendar popup"""
        top = tk.Toplevel(self.root)
        top.title("Select Date")
        cal = Calendar(top, selectmode='day', date_pattern='y-mm-dd')
        cal.pack(padx=10, pady=10)
        ttk.Button(top, text="Select", command=lambda: self.set_date(cal.get_date(), top)).pack(pady=5)

    def set_date(self, selected_date, top):
        """Set selected date from calendar"""
        self.date_entry.delete(0, 'end')
        self.date_entry.insert(0, selected_date)
        top.destroy()
        self.update_weekday()

    def update_weekday(self):
        """Update weekday based on selected date"""
        try:
            selected_date = datetime.strptime(self.date_entry.get(), "%Y-%m-%d")
            weekday = selected_date.strftime("%A")
            if weekday in ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday"]:
                self.day_combo.set(weekday)
                self.update_points()
        except:
            pass

    def update_test_ui(self, event=None):
        """Update UI based on test type"""
        test_type = self.test_type.get()
        
        if test_type == "Daily":
            self.day_label.grid(row=0, column=5, padx=5, pady=5, sticky='e')
            self.day_combo.grid(row=0, column=6, padx=5, pady=5, sticky='w')
            self.day_combo.bind("<<ComboboxSelected>>", self.update_points)
            
            # Show additional point controls for daily tests
            self.add_point_label.grid(row=0, column=7, padx=5, pady=5, sticky='e')
            self.add_point_combo.grid(row=0, column=8, padx=5, pady=5, sticky='w')
            self.add_point_btn.grid(row=0, column=9, padx=5, pady=5)
            
            # Update available points for additional selection
            day = self.day_combo.get()
            if day:
                # Combine used points from both micro and chem
                used_points = set(self.DAILY_MICRO_POINTS.get(day, []) + self.DAILY_CHEM_POINTS.get(day, []))
                available_points = [p for p in self.ALL_POINTS if p not in used_points]
                self.add_point_combo['values'] = available_points
                if available_points:
                    self.add_point_combo.current(0)
            
            self.update_weekday()
        else:
            self.day_label.grid_remove()
            self.day_combo.grid_remove()
            self.add_point_label.grid_remove()
            self.add_point_combo.grid_remove()
            self.add_point_btn.grid_remove()
            
        self.update_points()

    def add_additional_point(self):
        """Add an additional point to the daily test"""
        test_type = self.test_type.get()
        day = self.day_combo.get() if test_type == "Daily" else None
        point = self.add_point_combo.get()
        
        if not point:
            return
            
        # Add to microbiology table if not already present
        micro_points = {self.micro_table.item(item)['values'][0] for item in self.micro_table.get_children()}
        if point not in micro_points:
            self.micro_table.insert("", 'end', values=(point, "", "Absent", "Absent", ""))
        
        # Add to chemistry table if not already present
        chem_points = {self.chem_table.item(item)['values'][0] for item in self.chem_table.get_children()}
        if point not in chem_points:
            self.chem_table.insert("", 'end', values=(point, "", "", "", ""))
        
        # Update available points
        used_points = set(self.DAILY_MICRO_POINTS.get(day, []) + 
                         self.DAILY_CHEM_POINTS.get(day, []) + 
                         list(micro_points))
        available_points = [p for p in self.ALL_POINTS if p not in used_points]
        self.add_point_combo['values'] = available_points
        if available_points:
            self.add_point_combo.current(0)
        else:
            self.add_point_combo.set('')

    def update_points(self, event=None):
        """Update points in both tabs"""
        test_type = self.test_type.get()
        day = self.day_combo.get() if test_type == "Daily" else None
        
        if test_type == "After Sanitization":
            points = self.ALL_POINTS
            micro_points = points
            chem_points = points
        elif test_type == "Monthly":
            micro_points = self.MONTHLY_MICRO
            chem_points = self.MONTHLY_CHEM
        elif test_type == "Daily" and day:
            micro_points = self.DAILY_MICRO_POINTS.get(day, [])
            chem_points = self.DAILY_CHEM_POINTS.get(day, [])
        else:
            micro_points = []
            chem_points = []
        
        # Update microbiology table
        self.micro_table.delete(*self.micro_table.get_children())
        for point in micro_points:
            self.micro_table.insert("", 'end', values=(point, "", "Absent", "Absent", ""))
        
        # Update chemistry table
        self.chem_table.delete(*self.chem_table.get_children())
        for point in chem_points:
            self.chem_table.insert("", 'end', values=(point, "", "", "", ""))
        
        # Update available additional points for daily tests
        if test_type == "Daily" and day:
            used_points = set(self.DAILY_MICRO_POINTS.get(day, []) + self.DAILY_CHEM_POINTS.get(day, []))
            available_points = [p for p in self.ALL_POINTS if p not in used_points]
            self.add_point_combo['values'] = available_points
            if available_points:
                self.add_point_combo.current(0)
            else:
                self.add_point_combo.set('')

    def setup_micro_tab(self):
        """Microbiology tab setup"""
        self.micro_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.micro_frame, text="Microbiology")
        
        # Create paned window for resizable sections
        paned_window = ttk.PanedWindow(self.micro_frame, orient='vertical')
        paned_window.pack(fill='both', expand=True)
        
        # Table frame
        table_frame = ttk.Frame(paned_window)
        paned_window.add(table_frame, weight=3)  # 70% of space
        
        # Table with scrollbars
        tree_scroll = ttk.Frame(table_frame)
        tree_scroll.pack(fill='both', expand=True)
        
        # Vertical scrollbar
        y_scroll = ttk.Scrollbar(tree_scroll)
        y_scroll.pack(side='right', fill='y')
        
        # Horizontal scrollbar
        x_scroll = ttk.Scrollbar(tree_scroll, orient='horizontal')
        x_scroll.pack(side='bottom', fill='x')
        
        # Table
        self.micro_table = ttk.Treeview(tree_scroll, 
                                      columns=("Point", "Total Count", "Coliforms", "Pseudomonas", "Status"), 
                                      show='headings',
                                      yscrollcommand=y_scroll.set,
                                      xscrollcommand=x_scroll.set)
        
        # Configure columns
        self.micro_table.heading("Point", text="Point")
        self.micro_table.heading("Total Count", text="Total Count (CFU/mL)")
        self.micro_table.heading("Coliforms", text="Coliforms")
        self.micro_table.heading("Pseudomonas", text="Pseudomonas")
        self.micro_table.heading("Status", text="Status")
        
        self.micro_table.column("Point", width=150, anchor='w')
        self.micro_table.column("Total Count", width=120, anchor='center')
        self.micro_table.column("Coliforms", width=100, anchor='center')
        self.micro_table.column("Pseudomonas", width=100, anchor='center')
        self.micro_table.column("Status", width=120, anchor='center')
        
        self.micro_table.pack(fill='both', expand=True)
        
        # Configure scrollbars
        y_scroll.config(command=self.micro_table.yview)
        x_scroll.config(command=self.micro_table.xview)
        
        # Configure colors
        self.micro_table.tag_configure('Non-Conform', background='#ffcccc')
        self.micro_table.tag_configure('Warning', background='#fff3cd')
        self.micro_table.tag_configure('Conform', background='#ccffcc')
        
        # Data entry frame
        entry_frame = ttk.Frame(paned_window)
        paned_window.add(entry_frame, weight=1)  # 30% of space
        
        # Point selection
        point_frame = ttk.Frame(entry_frame)
        point_frame.pack(fill='x', padx=5, pady=5)
        ttk.Label(point_frame, text="Selected Point:").pack(side='left', padx=5)
        self.current_micro_point = ttk.Label(point_frame, text="None", width=20)
        self.current_micro_point.pack(side='left', padx=5)
        
        # Data entry fields
        data_frame = ttk.Frame(entry_frame)
        data_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(data_frame, text="Total Count:").grid(row=0, column=0, padx=5, sticky='e')
        self.micro_count = ttk.Entry(data_frame, width=10)
        self.micro_count.grid(row=0, column=1, padx=5, sticky='w')
        
        ttk.Label(data_frame, text="Coliforms:").grid(row=1, column=0, padx=5, sticky='e')
        self.coliforms = ttk.Combobox(data_frame, values=["Absent", "Present"], state='readonly', width=10)
        self.coliforms.set("Absent")
        self.coliforms.grid(row=1, column=1, padx=5, sticky='w')
        
        ttk.Label(data_frame, text="Pseudomonas:").grid(row=2, column=0, padx=5, sticky='e')
        self.pseudomonas = ttk.Combobox(data_frame, values=["Absent", "Present"], state='readonly', width=10)
        self.pseudomonas.set("Absent")
        self.pseudomonas.grid(row=2, column=1, padx=5, sticky='w')
        
        # Add button
        button_frame = ttk.Frame(entry_frame)
        button_frame.pack(fill='x', padx=5, pady=5)
        ttk.Button(button_frame, text="Add Data", command=self.add_micro_data).pack(pady=5)
        
        # Comments
        comments_frame = ttk.Frame(entry_frame)
        comments_frame.pack(fill='both', expand=True, padx=5, pady=5)
        ttk.Label(comments_frame, text="Comments:").pack(anchor='w')
        self.micro_comments = tk.Text(comments_frame, height=5, wrap='word')
        self.micro_comments.pack(fill='both', expand=True)
        
        # Bind selection
        self.micro_table.bind("<<TreeviewSelect>>", lambda e: self.on_table_select("micro"))

    def setup_chem_tab(self):
        """Chemistry tab setup"""
        self.chem_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.chem_frame, text="Chemistry")
        
        # Create paned window for resizable sections
        paned_window = ttk.PanedWindow(self.chem_frame, orient='vertical')
        paned_window.pack(fill='both', expand=True)
        
        # Table frame
        table_frame = ttk.Frame(paned_window)
        paned_window.add(table_frame, weight=3)  # 70% of space
        
        # Table with scrollbars
        tree_scroll = ttk.Frame(table_frame)
        tree_scroll.pack(fill='both', expand=True)
        
        # Vertical scrollbar
        y_scroll = ttk.Scrollbar(tree_scroll)
        y_scroll.pack(side='right', fill='y')
        
        # Horizontal scrollbar
        x_scroll = ttk.Scrollbar(tree_scroll, orient='horizontal')
        x_scroll.pack(side='bottom', fill='x')
        
        # Table
        self.chem_table = ttk.Treeview(tree_scroll, 
                                     columns=("Point", "Conductivity", "Oxidizable", "Cl Test", "Status"), 
                                     show='headings',
                                     yscrollcommand=y_scroll.set,
                                     xscrollcommand=x_scroll.set)
        
        # Configure columns
        self.chem_table.heading("Point", text="Point")
        self.chem_table.heading("Conductivity", text="Conductivity (µS/cm)")
        self.chem_table.heading("Oxidizable", text="Oxidizable Substances")
        self.chem_table.heading("Cl Test", text="Chloride Test (ppm)")
        self.chem_table.heading("Status", text="Status")
        
        self.chem_table.column("Point", width=150, anchor='w')
        self.chem_table.column("Conductivity", width=120, anchor='center')
        self.chem_table.column("Oxidizable", width=150, anchor='center')
        self.chem_table.column("Cl Test", width=120, anchor='center')
        self.chem_table.column("Status", width=120, anchor='center')
        
        self.chem_table.pack(fill='both', expand=True)
        
        # Configure scrollbars
        y_scroll.config(command=self.chem_table.yview)
        x_scroll.config(command=self.chem_table.xview)
        
        # Configure colors
        self.chem_table.tag_configure('Non-Conform', background='#ffcccc')
        self.chem_table.tag_configure('Warning', background='#fff3cd')
        self.chem_table.tag_configure('Conform', background='#ccffcc')
        
        # Data entry frame
        entry_frame = ttk.Frame(paned_window)
        paned_window.add(entry_frame, weight=1)  # 30% of space
        
        # Point selection
        point_frame = ttk.Frame(entry_frame)
        point_frame.pack(fill='x', padx=5, pady=5)
        ttk.Label(point_frame, text="Selected Point:").pack(side='left', padx=5)
        self.current_chem_point = ttk.Label(point_frame, text="None", width=20)
        self.current_chem_point.pack(side='left', padx=5)
        
        # Data entry fields
        data_frame = ttk.Frame(entry_frame)
        data_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(data_frame, text="Conductivity:").grid(row=0, column=0, padx=5, sticky='e')
        self.conductivity = ttk.Entry(data_frame, width=10)
        self.conductivity.grid(row=0, column=1, padx=5, sticky='w')
        
        ttk.Label(data_frame, text="Oxidizable:").grid(row=1, column=0, padx=5, sticky='e')
        self.oxidizable = ttk.Combobox(data_frame, values=["No color change", "Color change"], state='readonly', width=12)
        self.oxidizable.set("No color change")
        self.oxidizable.grid(row=1, column=1, padx=5, sticky='w')
        
        ttk.Label(data_frame, text="Cl Test:").grid(row=2, column=0, padx=5, sticky='e')
        self.cl_test = ttk.Entry(data_frame, width=10)
        self.cl_test.grid(row=2, column=1, padx=5, sticky='w')
        
        # Add button
        button_frame = ttk.Frame(entry_frame)
        button_frame.pack(fill='x', padx=5, pady=5)
        ttk.Button(button_frame, text="Add Data", command=self.add_chem_data).pack(pady=5)
        
        # Comments
        comments_frame = ttk.Frame(entry_frame)
        comments_frame.pack(fill='both', expand=True, padx=5, pady=5)
        ttk.Label(comments_frame, text="Comments:").pack(anchor='w')
        self.chem_comments = tk.Text(comments_frame, height=5, wrap='word')
        self.chem_comments.pack(fill='both', expand=True)
        
        # Bind selection
        self.chem_table.bind("<<TreeviewSelect>>", lambda e: self.on_table_select("chem"))

    def on_table_select(self, tab_type):
        """Update current point label when selection changes"""
        table = self.micro_table if tab_type == "micro" else self.chem_table
        label = self.current_micro_point if tab_type == "micro" else self.current_chem_point
        
        selected = table.focus()
        if selected:
            point = table.item(selected)["values"][0]
            label.config(text=point)

    def add_micro_data(self):
        """Add microbiology data with point-specific limits"""
        selected = self.micro_table.focus()
        if not selected:
            messagebox.showerror("Error", "Please select a point first!")
            return
        
        point = self.micro_table.item(selected)["values"][0]
        count = self.micro_count.get()
        coliforms = self.coliforms.get()
        pseudomonas = self.pseudomonas.get()
        comments = self.micro_comments.get("1.0", 'end-1c')
        
        # Get the specific limit for this point
        limit = self.CFU_LIMITS.get(point, 500)
        
        # Validation
        status = "Conform"
        if count.isdigit():
            count_int = int(count)
            if count_int > limit:
                status = "Non-Conform"
            elif count_int > 0.4 * limit:
                status = "Warning"
        else:
            status = "Invalid Input"
        
        if coliforms == "Present" or pseudomonas == "Present":
            status = "Non-Conform (Microbial)"
        
        # Update table
        self.micro_table.item(selected, 
                            values=(point, count, coliforms, pseudomonas, status),
                            tags=(status,))
        
        # Save data
        self.current_data.append({
            "Date": self.date_entry.get(),
            "Test Type": self.test_type.get(),
            "Day": self.day_combo.get() if self.test_type.get() == "Daily" else "",
            "Point": point,
            "Total Count": count,
            "Coliforms": coliforms,
            "Pseudomonas": pseudomonas,
            "Status": status,
            "Comments": comments,
            "Tab": "Microbiology"
        })
        
        # Clear entries
        self.micro_count.delete(0, 'end')
        self.coliforms.set("Absent")
        self.pseudomonas.set("Absent")
        self.micro_comments.delete("1.0", tk.END)

    def add_chem_data(self):
        """Add chemistry data to table with conformance checking"""
        selected = self.chem_table.focus()
        if not selected:
            messagebox.showerror("Error", "Please select a point first!")
            return
        
        point = self.chem_table.item(selected)["values"][0]
        conductivity = self.conductivity.get()
        oxidizable = self.oxidizable.get()
        cl_test = self.cl_test.get()
        comments = self.chem_comments.get("1.0", 'end-1c')
        
        # Check conformance
        status = "Conform"
        issues = []
        
        # Check conductivity
        if conductivity:
            try:
                cond_value = float(conductivity)
                limit = self.CHEM_LIMITS["Conductivity"].get(point, 1.3)  # Default PW limit
                if cond_value > limit:
                    issues.append(f"Conductivity > {limit} µS/cm")
            except ValueError:
                issues.append("Invalid conductivity value")
        
        # Check oxidizable substances
        if oxidizable == "Color change":
            issues.append("Oxidizable substances detected")
        
        # Check chloride test
        if cl_test:
            try:
                cl_value = float(cl_test)
                # Special points where >0.5 is acceptable (conform)
                if point in self.CHEM_LIMITS["Cl_Allowed_Points"]:
                    if cl_value <= 0.5:  # Only non-conform if ≤ 0.5 for these points
                        issues.append(f"Chloride ≤ 0.5 ppm (needs to be > 0.5 for this point)")
                else:
                    # For all other points, any chloride is non-conform
                    if cl_value > 0:
                        issues.append("Chloride detected (should be 0)")
            except ValueError:
                issues.append("Invalid chloride value")
        
        # Determine final status
        if issues:
            status = "Non-Conform: " + ", ".join(issues)
        elif not conductivity and not cl_test:
            status = "Incomplete data"
        
        # Update table
        self.chem_table.item(selected, 
                           values=(point, conductivity, oxidizable, cl_test, status),
                           tags=("Non-Conform" if issues else "Conform",))
        
        # Save data
        self.current_data.append({
            "Date": self.date_entry.get(),
            "Test Type": self.test_type.get(),
            "Day": self.day_combo.get() if self.test_type.get() == "Daily" else "",
            "Point": point,
            "Conductivity": conductivity,
            "Oxidizable": oxidizable,
            "Cl Test": cl_test,
            "Status": status,
            "Comments": comments,
            "Tab": "Chemistry"
        })
        
        # Clear entries
        self.conductivity.delete(0, 'end')
        self.oxidizable.set("No color change")
        self.cl_test.delete(0, 'end')
        self.chem_comments.delete("1.0", tk.END)

    def check_for_duplicates(self, df, new_record):
        """Check if a record with same date, point and test type already exists"""
        mask = (
            (df['Date'] == new_record['Date']) & 
            (df['Point'] == new_record['Point']) & 
            (df['Test Type'] == new_record['Test Type'])
        )
        return df[mask].empty

    def export_data(self):
        """Export all data to appropriate database files using pandas"""
        if not self.current_data:
            messagebox.showerror("Error", "No data to export!")
            return
        
        exported_files = set()
        errors = []
        
        for record in self.current_data:
            test_type = record["Test Type"]
            is_micro = record["Tab"] == "Microbiology"
            
            # Determine which database file to use
            if test_type == "After Sanitization":
                db_key = f"Sanitization_{'Micro' if is_micro else 'Chem'}"
            else:
                db_key = f"{test_type}_{'Micro' if is_micro else 'Chem'}"
                
            filename = self.DB_FILES[db_key]
            filepath = os.path.join("QC_Databases", filename)
            
            try:
                # Read existing data
                if os.path.exists(filepath):
                    df = pd.read_csv(filepath)
                else:
                    if is_micro:
                        df = pd.DataFrame(columns=[
                            "Date", "Test Type", "Day", "Point",
                            "Total Count", "Coliforms", "Pseudomonas",
                            "Status", "Comments"
                        ])
                    else:
                        df = pd.DataFrame(columns=[
                            "Date", "Test Type", "Day", "Point",
                            "Conductivity", "Oxidizable", "Cl Test",
                            "Status", "Comments"
                        ])
                
                # Check for duplicates
                if not self.check_for_duplicates(df, record):
                    errors.append(f"Duplicate entry for {record['Point']} on {record['Date']}")
                    continue
                
                # Prepare new row
                new_row = {
                    "Date": record["Date"],
                    "Test Type": test_type,
                    "Day": record.get("Day", ""),
                    "Point": record["Point"],
                    "Comments": record.get("Comments", ""),
                    "Status": record.get("Status", "")
                }
                
                if is_micro:
                    new_row.update({
                        "Total Count": record["Total Count"],
                        "Coliforms": record["Coliforms"],
                        "Pseudomonas": record["Pseudomonas"]
                    })
                else:
                    new_row.update({
                        "Conductivity": record.get("Conductivity", ""),
                        "Oxidizable": record.get("Oxidizable", ""),
                        "Cl Test": record.get("Cl Test", "")
                    })
                
                # Append new data
                df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
                
                # Save back to file
                df.to_csv(filepath, index=False)
                exported_files.add(filepath)
                
            except Exception as e:
                errors.append(f"Error saving {record['Point']}: {str(e)}")
        
        # Show results
        message_lines = []
        
        if exported_files:
            message_lines.append("Data successfully saved to:")
            message_lines.extend([f"- {os.path.abspath(f)}" for f in exported_files])
            self.status_label.config(text="\n".join(message_lines), foreground='green')
        
        if errors:
            message_lines.append("\nErrors encountered:")
            message_lines.extend([f"- {e}" for e in errors])
            self.status_label.config(text="\n".join(message_lines), foreground='red')
        
        if exported_files or errors:
            messagebox.showinfo("Export Complete", "\n".join(message_lines))
        
        # Clear current data only after successful export
        self.current_data = []

if __name__ == "__main__":
    root = tk.Tk()
    
    # Set Windows 10/11 theme if available
    if sys.platform == "win32":
        try:
            from ctypes import windll
            windll.shcore.SetProcessDpiAwareness(1)
        except:
            pass
    
    # Install required packages if missing
    try:
        from tkcalendar import Calendar
        import pandas as pd
        from docx import Document
        from matplotlib.figure import Figure
    except ImportError:
        import subprocess
        subprocess.run(["pip", "install", "tkcalendar", "pandas", "python-docx", "matplotlib"])
        from tkcalendar import Calendar
        import pandas as pd
        from docx import Document
        from matplotlib.figure import Figure
    
    app = WaterQCApp(root)
    root.mainloop()