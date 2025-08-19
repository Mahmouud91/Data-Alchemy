#!/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import csv
from datetime import datetime
import os

# Custom Entry with placeholder functionality
class PlaceholderEntry(ttk.Entry):
    def __init__(self, container, placeholder, *args, **kwargs):
        super().__init__(container, *args, style="Placeholder.TEntry", **kwargs)
        self.placeholder = placeholder
        self.insert("0", self.placeholder)
        self.bind("<FocusIn>", self._clear_placeholder)
        self.bind("<FocusOut>", self._add_placeholder)
        self.configure(style="TEntry")

    def _clear_placeholder(self, e):
        if self.get() == self.placeholder:
            self.delete("0", "end")
            self.configure(style="TEntry")

    def _add_placeholder(self, e):
        if not self.get():
            self.insert("0", self.placeholder)
            self.configure(style="Placeholder.TEntry")

    def get(self):
        content = super().get()
        if content == self.placeholder:
            return ""
        return content

# === AQL Tables ===
aql_tables = {
    "Level 1": [
        {"min": 2, "max": 8, "sample": 2, "major": 0, "minor": 0},
        {"min": 9, "max": 15, "sample": 2, "major": 0, "minor": 0},
        {"min": 16, "max": 25, "sample": 3, "major": 0, "minor": 0},
        {"min": 26, "max": 50, "sample": 5, "major": 0, "minor": 0},
        {"min": 51, "max": 90, "sample": 5, "major": 0, "minor": 0},
        {"min": 91, "max": 150, "sample": 8, "major": 0, "minor": 1},
        {"min": 151, "max": 280, "sample": 13, "major": 0, "minor": 1},
        {"min": 281, "max": 500, "sample": 20, "major": 1, "minor": 2},
        {"min": 501, "max": 1200, "sample": 32, "major": 2, "minor": 3},
        {"min": 1201, "max": 3200, "sample": 50, "major": 3, "minor": 5},
        {"min": 3201, "max": 10000, "sample": 80, "major": 5, "minor": 7},
        {"min": 10001, "max": 35000, "sample": 125, "major": 7, "minor": 10},
        {"min": 35001, "max": 150000, "sample": 200, "major": 10, "minor": 14},
        {"min": 150001, "max": 500000, "sample": 315, "major": 14, "minor": 21},
        {"min": 500001, "max": float('inf'), "sample": 500, "major": 21, "minor": 21}
    ],
    "Level 2": [
        {"min": 2, "max": 8, "sample": 2, "major": 0, "minor": 0},
        {"min": 9, "max": 15, "sample": 3, "major": 0, "minor": 0},
        {"min": 16, "max": 25, "sample": 5, "major": 0, "minor": 0},
        {"min": 26, "max": 50, "sample": 8, "major": 0, "minor": 1},
        {"min": 51, "max": 90, "sample": 13, "major": 1, "minor": 1},
        {"min": 91, "max": 150, "sample": 20, "major": 1, "minor": 2},
        {"min": 151, "max": 280, "sample": 32, "major": 2, "minor": 3},
        {"min": 281, "max": 500, "sample": 50, "major": 3, "minor": 5},
        {"min": 501, "max": 1200, "sample": 80, "major": 5, "minor": 7},
        {"min": 1201, "max": 3200, "sample": 125, "major": 7, "minor": 10},
        {"min": 3201, "max": 10000, "sample": 200, "major": 10, "minor": 14},
        {"min": 10001, "max": 35000, "sample": 315, "major": 14, "minor": 21},
        {"min": 35001, "max": 150000, "sample": 500, "major": 21, "minor": 21},
        {"min": 150001, "max": 500000, "sample": 800, "major": 21, "minor": 21},
        {"min": 500001, "max": float('inf'), "sample": 1250, "major": 21, "minor": 21}
    ],
    "S-4": [
        {"min": 2, "max": 8, "sample": 2, "major": 0, "minor": 0},
        {"min": 9, "max": 15, "sample": 2, "major": 0, "minor": 0},
        {"min": 16, "max": 25, "sample": 3, "major": 0, "minor": 0},
        {"min": 26, "max": 50, "sample": 5, "major": 0, "minor": 0},
        {"min": 51, "max": 90, "sample": 5, "major": 0, "minor": 0},
        {"min": 91, "max": 150, "sample": 8, "major": 0, "minor": 1},
        {"min": 151, "max": 280, "sample": 13, "major": 0, "minor": 1},
        {"min": 281, "max": 500, "sample": 13, "major": 0, "minor": 1},
        {"min": 501, "max": 1200, "sample": 20, "major": 1, "minor": 2},
        {"min": 1201, "max": 3200, "sample": 32, "major": 2, "minor": 3},
        {"min": 3201, "max": 10000, "sample": 32, "major": 2, "minor": 3},
        {"min": 10001, "max": 35000, "sample": 50, "major": 3, "minor": 5},
        {"min": 35001, "max": 150000, "sample": 80, "major": 5, "minor": 7},
        {"min": 150001, "max": 500000, "sample": 80, "major": 5, "minor": 7},
        {"min": 500001, "max": float('inf'), "sample": 125, "major": 7, "minor": 10}
    ]
}

# === Raw Material Types and Tests ===
tests_by_type = {
    "bottle": ["Volume", "Length & Width", "Leakage Test", "Appearance"],
    "Cap": ["Length & Width", "Leakage Test", "Appearance"],
    "alu pouch": ["Appearance and Color", "Length & Width", "Leakage Test"],
    "plastic cassette": ["Appearance", "Length & Width", "Compatibility Between Top and Bottom"],
    "silica gel": ["Length & Width", "Weight"],
    "uncut sheet": ["Length & Width", "Controls & Flow Rate", "Compatibility with Cassette"],
    "soft bag": ["Volume", "Leakage"],
    "cbc carton": ["Volume", "Length & Width & Height", "Visual Appearance", "Stacking"],
    "lateral flow box": ["Length & Width & Height", "Visual Appearance"],
    "stickers": ["Length & Width", "Check data,codes (product code & sticker code)", "Visual Appearance"]
}

class AQLInspector:
    def __init__(self, root):
        self.root = root
        self.root.title("AQL Inspection System")
        self.root.geometry("1100x800")

        # Create style for placeholder
        self.style = ttk.Style()
        self.style.configure("Placeholder.TEntry", foreground="grey")

        self.setup_ui()

    def setup_ui(self):
        # Notebook for multiple tabs
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        # Tab 1: Inspection Plan
        self.setup_inspection_tab()

        # Tab 2: Conformity Check
        self.setup_conformity_tab()

        # Tab 3: Search Records
        self.setup_search_tab()

    def setup_inspection_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Inspection Plan")

        # Input Frame
        input_frame = ttk.LabelFrame(tab, text="Inspection Details", padding=10)
        input_frame.pack(fill=tk.X, padx=10, pady=5)

        # Row 0
        ttk.Label(input_frame, text="Internal Code:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
        self.ic_entry = ttk.Entry(input_frame, width=30)
        self.ic_entry.grid(row=0, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Product Name:").grid(row=0, column=2, sticky="e", padx=5, pady=5)
        self.product_name_entry = ttk.Entry(input_frame, width=30)
        self.product_name_entry.grid(row=0, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Product Code:").grid(row=0, column=4, sticky="e", padx=5, pady=5)
        self.product_code_entry = ttk.Entry(input_frame, width=30)
        self.product_code_entry.grid(row=0, column=5, sticky="w", padx=5, pady=5)

        # Row 1
        ttk.Label(input_frame, text="Sampler Name:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
        self.sampler_entry = ttk.Entry(input_frame, width=30)
        self.sampler_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Supplier Name:").grid(row=1, column=2, sticky="e", padx=5, pady=5)
        self.supplier_entry = ttk.Entry(input_frame, width=30)
        self.supplier_entry.grid(row=1, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Number of Units:").grid(row=1, column=4, sticky="e", padx=5, pady=5)
        self.units_entry = ttk.Entry(input_frame, width=30)
        self.units_entry.grid(row=1, column=5, sticky="w", padx=5, pady=5)

        # Row 2
        ttk.Label(input_frame, text="Item Type:").grid(row=2, column=0, sticky="e", padx=5, pady=5)
        self.item_var = tk.StringVar()
        item_menu = ttk.Combobox(input_frame, textvariable=self.item_var, 
                                values=list(tests_by_type.keys()), state="readonly", width=27)
        item_menu.grid(row=2, column=1, sticky="w", padx=5, pady=5)
        item_menu.current(0)

        ttk.Label(input_frame, text="Inspection Level:").grid(row=2, column=2, sticky="e", padx=5, pady=5)
        self.level_var = tk.StringVar()
        level_menu = ttk.Combobox(input_frame, textvariable=self.level_var, 
                                 values=["Level 1", "Level 2", "S-4"], state="readonly", width=27)
        level_menu.grid(row=2, column=3, sticky="w", padx=5, pady=5)
        level_menu.current(1)

        # Button
        ttk.Button(tab, text="Generate Inspection Plan", 
                  command=self.generate_inspection_plan).pack(pady=10)

        # Output Frame
        output_frame = ttk.LabelFrame(tab, text="Inspection Plan", padding=10)
        output_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.output_text = scrolledtext.ScrolledText(output_frame, wrap=tk.WORD, 
                                                   width=80, height=20, font=('Courier', 10))
        self.output_text.pack(fill=tk.BOTH, expand=True)

    def setup_conformity_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Conformity Check")

        # Input Frame
        input_frame = ttk.LabelFrame(tab, text="Conformity Details", padding=10)
        input_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(input_frame, text="Internal Code:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
        self.conform_ic_entry = ttk.Entry(input_frame, width=30)
        self.conform_ic_entry.grid(row=0, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Product Name:").grid(row=0, column=2, sticky="e", padx=5, pady=5)
        self.conform_product_name_entry = ttk.Entry(input_frame, width=30)
        self.conform_product_name_entry.grid(row=0, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Product Code:").grid(row=0, column=4, sticky="e", padx=5, pady=5)
        self.conform_product_code_entry = ttk.Entry(input_frame, width=30)
        self.conform_product_code_entry.grid(row=0, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Inspector Name:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
        self.inspector_entry = ttk.Entry(input_frame, width=30)
        self.inspector_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Major Defects Found:").grid(row=1, column=2, sticky="e", padx=5, pady=5)
        self.major_defects_entry = ttk.Entry(input_frame, width=30)
        self.major_defects_entry.grid(row=1, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Minor Defects Found:").grid(row=1, column=4, sticky="e", padx=5, pady=5)
        self.minor_defects_entry = ttk.Entry(input_frame, width=30)
        self.minor_defects_entry.grid(row=1, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(input_frame, text="Status:").grid(row=2, column=0, sticky="e", padx=5, pady=5)
        self.status_var = tk.StringVar()
        status_menu = ttk.Combobox(input_frame, textvariable=self.status_var, 
                                  values=["Conform", "Non-Conform"], state="readonly", width=27)
        status_menu.grid(row=2, column=1, sticky="w", padx=5, pady=5)
        status_menu.current(0)

        ttk.Label(input_frame, text="Comments:").grid(row=3, column=0, sticky="ne", padx=5, pady=5)
        self.comments_entry = tk.Text(input_frame, width=70, height=4)
        self.comments_entry.grid(row=3, column=1, columnspan=5, sticky="we", padx=5, pady=5)

        # Buttons
        button_frame = ttk.Frame(tab)
        button_frame.pack(pady=10)

        ttk.Button(button_frame, text="Save Conformity & Generate Certificate", 
                  command=self.save_conformity).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Clear", 
                  command=self.clear_conformity).pack(side=tk.LEFT, padx=5)

    def setup_search_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Search Records")

        # Search Frame
        search_frame = ttk.LabelFrame(tab, text="Search Criteria", padding=10)
        search_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(search_frame, text="Search by IC:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
        self.search_ic_entry = ttk.Entry(search_frame, width=30)
        self.search_ic_entry.grid(row=0, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(search_frame, text="Product Name:").grid(row=0, column=2, sticky="e", padx=5, pady=5)
        self.search_product_name_entry = ttk.Entry(search_frame, width=30)
        self.search_product_name_entry.grid(row=0, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(search_frame, text="Date Range:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
        self.start_date_entry = PlaceholderEntry(search_frame, "YYYY-MM-DD", width=12)
        self.start_date_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(search_frame, text="to").grid(row=1, column=2, padx=5, pady=5)

        self.end_date_entry = PlaceholderEntry(search_frame, "YYYY-MM-DD", width=12)
        self.end_date_entry.grid(row=1, column=3, sticky="w", padx=5, pady=5)

        ttk.Button(search_frame, text="Search", 
                  command=self.search_records).grid(row=1, column=4, padx=5, pady=5)

        # Results Frame
        results_frame = ttk.LabelFrame(tab, text="Search Results", padding=10)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Treeview for results
        self.results_tree = ttk.Treeview(results_frame, columns=("Timestamp", "IC", "ProductName", "ProductCode", "Sampler", "Supplier", "Units", "Item", "Level", 
                                                               "Sample", "Major", "Minor", "Status", "Inspector"), 
                                       show="headings")

        # Define headings
        self.results_tree.heading("Timestamp", text="Timestamp")
        self.results_tree.heading("IC", text="IC")
        self.results_tree.heading("ProductName", text="Product Name")
        self.results_tree.heading("ProductCode", text="Product Code")
        self.results_tree.heading("Sampler", text="Sampler")
        self.results_tree.heading("Supplier", text="Supplier")
        self.results_tree.heading("Units", text="Units")
        self.results_tree.heading("Item", text="Item")
        self.results_tree.heading("Level", text="Level")
        self.results_tree.heading("Sample", text="Sample")
        self.results_tree.heading("Major", text="Major")
        self.results_tree.heading("Minor", text="Minor")
        self.results_tree.heading("Status", text="Status")
        self.results_tree.heading("Inspector", text="Inspector")

        # Set column widths
        self.results_tree.column("Timestamp", width=150)
        self.results_tree.column("IC", width=100)
        self.results_tree.column("ProductName", width=120)
        self.results_tree.column("ProductCode", width=100)
        self.results_tree.column("Sampler", width=100)
        self.results_tree.column("Supplier", width=120)
        self.results_tree.column("Units", width=60)
        self.results_tree.column("Item", width=120)
        self.results_tree.column("Level", width=80)
        self.results_tree.column("Sample", width=60)
        self.results_tree.column("Major", width=60)
        self.results_tree.column("Minor", width=60)
        self.results_tree.column("Status", width=100)
        self.results_tree.column("Inspector", width=120)

        # Add scrollbar
        scrollbar = ttk.Scrollbar(results_frame, orient="vertical", command=self.results_tree.yview)
        scrollbar.pack(side="right", fill="y")
        self.results_tree.configure(yscrollcommand=scrollbar.set)

        self.results_tree.pack(fill=tk.BOTH, expand=True)

        # Export button
        ttk.Button(tab, text="Export to CSV", 
                  command=self.export_results).pack(pady=10)

    def get_aql_values(self, units, level):
        for row in aql_tables[level]:
            if row["min"] <= units <= row["max"]:
                return row["sample"], row["major"], row["minor"]
        return None, None, None

    def generate_inspection_plan(self):
        ic = self.ic_entry.get().strip()
        product_name = self.product_name_entry.get().strip()
        product_code = self.product_code_entry.get().strip()
        sampler = self.sampler_entry.get().strip()
        supplier = self.supplier_entry.get().strip()

        if not ic:
            messagebox.showwarning("Warning", "Please enter an Internal Code")
            return
        if not sampler:
            messagebox.showwarning("Warning", "Please enter Sampler Name")
            return

        level = self.level_var.get()
        try:
            units = int(self.units_entry.get())
            if units <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid positive number of units.")
            return

        item = self.item_var.get()
        sample, major, minor = self.get_aql_values(units, level)
        if sample is None:
            messagebox.showerror("Error", "Units not covered in AQL table for this level.")
            return

        tests = tests_by_type[item]

        output = f"📋 Inspection Plan\n{'='*40}\n"
        output += f"• Internal Code: {ic}\n"
        output += f"• Product Name: {product_name}\n"
        output += f"• Product Code: {product_code}\n"
        output += f"• Sampler: {sampler}\n"
        output += f"• Supplier: {supplier}\n"
        output += f"• Units: {units}\n"
        output += f"• Item Type: {item}\n"
        output += f"• Inspection Level: {level}\n"
        output += f"• Sample Size: {sample}\n"
        output += f"• Major Defects (2.5%): Accept ≤ {major}, Reject ≥ {major+1}\n"
        output += f"• Minor Defects (4.0%): Accept ≤ {minor}, Reject ≥ {minor+1}\n"
        output += f"\n🔍 Required Tests:\n"
        for test in tests:
            output += f"  - {test}\n"
        output += f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"

        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, output)

        # Save to CSV
        self.save_to_csv(ic, product_name, product_code, sampler, supplier, units, item, level, sample, tests, major, minor)

    def save_to_csv(self, ic, product_name, product_code, sampler, supplier, units, item, level, sample, tests, major, minor):
        file_exists = os.path.isfile("inspection_results.csv")
        with open("inspection_results.csv", mode="a", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            if not file_exists:
                writer.writerow([
                    "Timestamp", "Internal Code", "Product Name", "Product Code", "Sampler", "Supplier", "Units", "Item Type", 
                    "Inspection Level", "Sample Size", "Required Tests",
                    "Major Defects", "Minor Defects", "Status", "Inspector", "Comments"
                ])
            writer.writerow([
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                ic, product_name, product_code, sampler, supplier, units, item, level, sample,
                ", ".join(tests),
                f"AQL2.5% Major: Ac {major}/Re {major + 1}",
                f"AQL4.0% Minor: Ac {minor}/Re {minor + 1}",
                "", "", ""  # Empty fields for conformity data
            ])
        messagebox.showinfo("Saved", "✅ Data saved to inspection_results.csv")

    def save_conformity(self):
        ic = self.conform_ic_entry.get().strip()
        product_name = self.conform_product_name_entry.get().strip()
        product_code = self.conform_product_code_entry.get().strip()
        inspector = self.inspector_entry.get().strip()
        status = self.status_var.get()
        comments = self.comments_entry.get("1.0", tk.END).strip()

        if not ic:
            messagebox.showwarning("Warning", "Please enter an Internal Code")
            return
        if not inspector:
            messagebox.showwarning("Warning", "Please enter Inspector Name")
            return

        try:
            major_defects = int(self.major_defects_entry.get())
        except ValueError:
            major_defects = 0

        try:
            minor_defects = int(self.minor_defects_entry.get())
        except ValueError:
            minor_defects = 0

        # Get additional info from CSV for the certificate
        supplier = ""
        item_type = ""
        units = ""
        sample_size = ""
        
        try:
            with open("inspection_results.csv", mode="r", newline="", encoding="utf-8") as file:
                reader = csv.DictReader(file)
                for row in reader:
                    if row["Internal Code"] == ic:
                        supplier = row.get("Supplier", "N/A")
                        item_type = row.get("Item Type", "N/A")
                        units = row.get("Units", "N/A")
                        sample_size = row.get("Sample Size", "N/A")
                        break
        except FileNotFoundError:
            pass

        # Update the CSV record
        updated = False
        rows = []

        try:
            with open("inspection_results.csv", mode="r", newline="", encoding="utf-8") as file:
                reader = csv.reader(file)
                headers = next(reader)
                rows.append(headers)

                for row in reader:
                    if row[1] == ic:  # Match by Internal Code
                        row[13] = status
                        row[14] = inspector
                        row[15] = comments
                        row[11] = f"Major Defects Found: {major_defects}"
                        row[12] = f"Minor Defects Found: {minor_defects}"
                        
                        # Update product name and code if they were provided
                        if product_name:
                            row[2] = product_name
                        if product_code:
                            row[3] = product_code
                            
                        updated = True
                    rows.append(row)

            if updated:
                with open("inspection_results.csv", mode="w", newline="", encoding="utf-8") as file:
                    writer = csv.writer(file)
                    writer.writerows(rows)
                
                # Generate Word document
                self.generate_certificate(
                    ic, product_name, product_code, supplier, item_type, units, sample_size, 
                    status, major_defects, minor_defects, inspector, comments
                )
                
                messagebox.showinfo("Success", "Conformity data updated and certificate generated successfully")
            else:
                messagebox.showerror("Error", f"No record found with IC: {ic}")

        except FileNotFoundError:
            messagebox.showerror("Error", "No inspection records found")

    def generate_certificate(self, ic, product_name, product_code, supplier, item_type, units, sample_size, 
                           status, major_defects, minor_defects, inspector, comments):
        doc = Document()
        
        # Add title
        title = doc.add_heading('RAW MATERIAL INSPECTION CERTIFICATE', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company info
        company = doc.add_paragraph()
        company.add_run("Company Name: ").bold = True
        company.add_run("Your Company Name Here\n")
        company.add_run("Address: ").bold = True
        company.add_run("123 Company Address, City, Country\n")
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add certificate number and date
        cert_info = doc.add_paragraph()
        cert_info.add_run(f"Certificate No: RM-{ic}-{datetime.now().strftime('%Y%m%d')}\n")
        cert_info.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
        cert_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        doc.add_paragraph("_"*50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add basic info table
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Shading Accent 1'
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        # Fill table
        data = [
            ("Internal Code:", ic),
            ("Product Name:", product_name),
            ("Product Code:", product_code),
            ("Supplier:", supplier),
            ("Material Type:", item_type),
            ("Batch Quantity:", units),
            ("Sample Size:", sample_size),
            ("Inspection Date:", datetime.now().strftime("%Y-%m-%d"))
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # Add inspection results
        doc.add_heading('Inspection Results', level=1)
        
        results_table = doc.add_table(rows=4, cols=2)
        results_data = [
            ("Status:", status),
            ("Major Defects Found:", str(major_defects)),
            ("Minor Defects Found:", str(minor_defects)),
            ("AQL Level:", "Level II (General Inspection Level)")
        ]
        
        for i, (label, value) in enumerate(results_data):
            results_table.cell(i, 0).text = label
            results_table.cell(i, 1).text = value
        
        # Add comments
        doc.add_heading('Comments', level=1)
        doc.add_paragraph(comments)
        
        # Add approval section
        doc.add_heading('Approval', level=1)
        approval_table = doc.add_table(rows=2, cols=2)
        approval_table.cell(0, 0).text = "Inspector:"
        approval_table.cell(0, 1).text = inspector
        approval_table.cell(1, 0).text = "Date:"
        approval_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d')
        
        # Add footer
        doc.add_paragraph("\n\n")
        footer = doc.add_paragraph()
        footer.add_run("This certificate is generated based on AQL inspection results.").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        if not os.path.exists("Certificates"):
            os.makedirs("Certificates")
            
        filename = f"Certificates/RM_Certificate_{ic}_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(filename)
        
        return filename

    def clear_conformity(self):
        self.conform_ic_entry.delete(0, tk.END)
        self.conform_product_name_entry.delete(0, tk.END)
        self.conform_product_code_entry.delete(0, tk.END)
        self.inspector_entry.delete(0, tk.END)
        self.major_defects_entry.delete(0, tk.END)
        self.minor_defects_entry.delete(0, tk.END)
        self.status_var.set("Conform")
        self.comments_entry.delete("1.0", tk.END)

    def search_records(self):
        search_ic = self.search_ic_entry.get().strip()
        search_product_name = self.search_product_name_entry.get().strip()
        start_date = self.start_date_entry.get()
        end_date = self.end_date_entry.get()

        # Clear previous results
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)

        try:
            with open("inspection_results.csv", mode="r", newline="", encoding="utf-8") as file:
                reader = csv.DictReader(file)

                for row in reader:
                    # Skip header if present
                    if "Timestamp" not in row:
                        continue

                    # Apply filters
                    match = True

                    # IC filter
                    if search_ic and search_ic.lower() not in row["Internal Code"].lower():
                        match = False

                    # Product name filter
                    if search_product_name and search_product_name.lower() not in row.get("Product Name", "").lower():
                        match = False

                    # Date range filter
                    if start_date or end_date:
                        try:
                            record_date = datetime.strptime(row["Timestamp"], "%Y-%m-%d %H:%M:%S").date()

                            if start_date:
                                start = datetime.strptime(start_date, "%Y-%m-%d").date()
                                if record_date < start:
                                    match = False

                            if end_date:
                                end = datetime.strptime(end_date, "%Y-%m-%d").date()
                                if record_date > end:
                                    match = False
                        except ValueError:
                            pass

                    if match:
                        self.results_tree.insert("", "end", values=(
                            row["Timestamp"],
                            row["Internal Code"],
                            row.get("Product Name", ""),
                            row.get("Product Code", ""),
                            row["Sampler"],
                            row.get("Supplier", ""),
                            row["Units"],
                            row["Item Type"],
                            row["Inspection Level"],
                            row["Sample Size"],
                            row["Major Defects"],
                            row["Minor Defects"],
                            row.get("Status", ""),
                            row.get("Inspector", "")
                        ))

        except FileNotFoundError:
            messagebox.showerror("Error", "No inspection records found")

    def export_results(self):
        items = self.results_tree.get_children()
        if not items:
            messagebox.showwarning("Warning", "No results to export")
            return

        try:
            with open("search_results.csv", mode="w", newline="", encoding="utf-8") as file:
                writer = csv.writer(file)

                # Write headers
                headers = [self.results_tree.heading(col)["text"] for col in self.results_tree["columns"]]
                writer.writerow(headers)

                # Write data
                for item in items:
                    row = self.results_tree.item(item)["values"]
                    writer.writerow(row)

            messagebox.showinfo("Success", f"Exported {len(items)} records to search_results.csv")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to export: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = AQLInspector(root)
    root.mainloop()